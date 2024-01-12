import docx as dx
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# implement RMA aouto compiler
# TODO TEMPLATE class - DONE


class Template:
    fileName = []

    def __init__(self, height=None, width=None, margins=None, header=None):
        self.height = height
        self.width = width
        self.margin = margins
        self.header = header
        # initialize worksheet
        self.document = Document()
        self.section = self.document.sections[0]
        self.section.page_height = Mm(self.height)
        self.section.page_width = Mm(self.width)
        self.section.left_margin = Mm(self.margin)
        self.section.right_margin = Mm(self.margin)
        self.section.top_margin = Mm(self.margin)
        self.section.bottom_margin = Mm(self.margin)
        self.section.header_distance = Mm(self.header)
        self.section.footer_distance = Mm(self.header)

    # create title
    def heading(self, title=str, level=int):
        self.title = title
        self.level = level
        self.document.add_heading(title, level)
        return None

    # create text
    def run(self, text=None, r=None, g=None, b=None, font=None, size=int):
        self.text = text
        self.r = r
        self.g = g
        self.b = b
        self.font = font
        self.size = size
        # define text, font, size, color
        run = self.document.add_paragraph().add_run(self.text)
        font = run.font
        color = font.color
        color.rgb = RGBColor(self.r, self.g, self.b)
        font.name = self.font
        font.size = Pt(self.size)
        # define spacing
        par_format = self.document.styles["Normal"].paragraph_format
        par_format.space_before = Pt(10)
        par_format.space_after = Pt(10)
        return None

    # create table
    def init_table(self, rw=int, cl=int, titles=str, size=int):
        self.rw = rw
        self.cl = cl
        self.table = self.document.add_table(rows=self.rw, cols=self.cl)
        # init columns titles
        self.row = self.table.rows[0].cells
        for i in range(len(titles)):
            self.row[i].text = titles[i]
        # define cell paragraph run font size
        for cell in self.row:
            par = cell.paragraphs[0]
            run = par.runs
            font = run[0].font
            font.size = Pt(size)
        return self.table

    # populate table
    def table_content(self, table=None, records=tuple, size=int):
        self.records = records
        self.table = table
        self.table.style = "Colorful List Accent 4"
        for a in range(len(self.records)):
            rowCells = self.table.add_row().cells
            for j in range(len(self.records[a])):
                if self.records[a][j]:
                    rowCells[j].text = str(self.records[a][j])
                else:
                    rowCells[j].text = input(f"inser value for:  {self.records[a][0]} ")
                    if "INV" in rowCells[j].text:
                        Template.fileName.append(rowCells[j].text)
            for cell in rowCells:
                par = cell.paragraphs[0]
                run = par.runs
                font = run[0].font
                font.size = Pt(size)
        return None

    # add new page
    def nextPage(self):
        self.document.add_page_break()

    # add image
    def addImg(self, path=str, size=int, align=None):
        """
        path = img path or name
        size = inces 1. to x.
        align = obj RIGHT or CENTER or LEFT
        """
        self.document.add_picture(path, width=Inches(size))
        par = self.document.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHTo
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHTo

    # save the document
    def save(self):
        name = Template.fileName[0]
        self.document.save(name + ".docx")
        return None
